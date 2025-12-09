"""
Text extraction for 10-Minute Reading App (V1).

Supports .docx and .txt files only.
Other formats are logged and skipped.
"""

from pathlib import Path
from typing import Optional


class ExtractionError(Exception):
    """Raised when text extraction fails."""
    pass


def extract_text(file_path: Path) -> str:
    """
    Extract text from a document.
    
    Args:
        file_path: Path to the document
    
    Returns:
        Extracted text content
    
    Raises:
        ExtractionError: If extraction fails or format unsupported
    """
    if not file_path.exists():
        raise ExtractionError(f"File not found: {file_path}")
    
    suffix = file_path.suffix.lower()
    
    if suffix == ".txt":
        return _extract_txt(file_path)
    elif suffix == ".docx":
        return _extract_docx(file_path)
    else:
        raise ExtractionError(f"Unsupported format: {suffix} (V1 only supports .txt and .docx)")


def _extract_txt(file_path: Path) -> str:
    """Extract text from a plain text file."""
    # Try common encodings
    encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
    
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read()
            return _clean_text(text)
        except UnicodeDecodeError:
            continue
    
    raise ExtractionError(f"Could not decode {file_path.name}")


def _extract_docx(file_path: Path) -> str:
    """Extract text from a Word document."""
    try:
        from docx import Document
    except ImportError:
        raise ExtractionError(
            "python-docx not installed. Run: pip install python-docx"
        )
    
    try:
        doc = Document(file_path)
        paragraphs = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        
        full_text = "\n\n".join(paragraphs)
        return _clean_text(full_text)
    
    except Exception as e:
        raise ExtractionError(f"Failed to read Word document: {e}")


def _clean_text(text: str) -> str:
    """Clean up extracted text."""
    if not text:
        return ""
    
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    
    # Remove excessive blank lines
    lines = text.split("\n")
    cleaned = []
    blank_count = 0
    
    for line in lines:
        line = line.strip()
        if not line:
            blank_count += 1
            if blank_count <= 1:
                cleaned.append("")
        else:
            blank_count = 0
            cleaned.append(line)
    
    return "\n".join(cleaned).strip()


def get_word_count(text: str) -> int:
    """Count words in text."""
    if not text:
        return 0
    return len(text.split())


def get_sentence_count(text: str) -> int:
    """Approximate sentence count."""
    if not text:
        return 0
    import re
    sentences = re.split(r"[.!?]+", text)
    return len([s for s in sentences if s.strip()])
